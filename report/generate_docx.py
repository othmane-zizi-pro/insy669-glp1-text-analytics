"""Generate a concise 2-page DOCX report for GLP-1 hybrid analysis."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

ROOT = Path(__file__).resolve().parent.parent
FIG  = ROOT / "figures"
DATA = ROOT / "data"

stats = json.loads((DATA / "analysis_stats.json").read_text())

doc = Document()

# ── Page setup: narrow margins ────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ── Styles ────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

for lvl, sz in [('Heading 1', 13), ('Heading 2', 11)]:
    s = doc.styles[lvl]
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(2)

# ── Helpers ───────────────────────────────────────────────
def add_fig(name, width=4.0, caption=None):
    p = FIG / name
    if p.exists():
        doc.add_picture(str(p), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].italic = True
            c.runs[0].font.size = Pt(8)
            c.paragraph_format.space_after = Pt(2)

def add_two_figs(n1, n2, c1=None, c2=None, width=2.4):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (nm, cap) in enumerate([(n1, c1), (n2, c2)]):
        cell = tbl.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp = FIG / nm
        if fp.exists():
            p.add_run().add_picture(str(fp), width=Inches(width))
        if cap:
            cp = cell.add_paragraph(cap)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(8)

def add_table(headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Shading Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8.5)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = tbl.rows[r+1].cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
run = title.add_run("Media vs. Public Opinion on GLP-1 Weight-Loss Drugs")
run.bold = True
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(1)
sub.add_run("A Comparative Text-Analytics Study of Ozempic & Wegovy").font.size = Pt(10)

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.paragraph_format.space_after = Pt(1)
r = auth.add_run(
    "V. Christopoulos, H. Guideau, S. Khosla, M. Yousuf, O. Zizi  \u2014  "
    "INSY 669 | McGill University | Winter 2026"
)
r.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "GLP-1 receptor agonists\u2014Ozempic and Wegovy\u2014have sparked extensive "
    "online patient discussion and media coverage. This study compares public discourse "
    "(3,246 Reddit posts + 102 WebMD reviews) with media discourse (634 Google News "
    "articles) from Jan\u2013Nov 2024, asking: (1) How does sentiment differ? "
    "(2) What linguistic/thematic gaps exist? (3) Does either stream lead temporally?"
)

# ══════════════════════════════════════════════════════════
# 2. DATA & METHODOLOGY
# ══════════════════════════════════════════════════════════
doc.add_heading("2. Data & Methodology", level=1)

add_table(
    ["Source", "Type", "n", "Method"],
    [
        ["Reddit", "Public", "3,246", "Arctic Shift API"],
        ["WebMD", "Public", "102", "Web scraping"],
        ["Google News", "Media", "634", "RSS + URL decode"],
    ]
)

doc.add_paragraph(
    "\nWe used a hybrid media-text strategy (full article body when available, snippet "
    "fallback otherwise). Preprocessing: tokenisation, stopword removal, lemmatisation. "
    "A length-normalised track (first 40 tokens) controls for document-length confounds. "
    "All ML pipelines use stratified k-fold CV to prevent leakage."
)

# ══════════════════════════════════════════════════════════
# 3. RESULTS
# ══════════════════════════════════════════════════════════
doc.add_heading("3. Results", level=1)

doc.add_heading("3.1 Sentiment Analysis", level=2)
doc.add_paragraph(
    f"VADER scores reveal a significant gap: public mean = +0.121, media = \u22120.137 "
    f"(t = 9.43, p < 0.001, Cohen\u2019s d = 0.44). Reddit is moderately positive "
    f"(+0.135), WebMD strongly negative (\u22120.308), news negative (\u22120.137). "
    f"All pairwise Mann\u2013Whitney U tests significant (p < 0.001)."
)
add_two_figs("sentiment_boxplot_3source.png", "sentiment_pies.png",
             "Sentiment by source", "Label distribution", width=2.4)

doc.add_heading("3.2 Corpus Comparison & Side Effects", level=2)
doc.add_paragraph(
    "TF-IDF confirms thematic divergence: public emphasises week, dose, lb, anyone; "
    "media foregrounds semaglutide, drug, weightloss drug. Cosine similarity: 0.346 "
    "(0.350 normalised). Public mentions nausea (3.38/1k tokens), diarrhea (1.22), "
    "constipation (1.16), anxiety (0.69) at rates far exceeding media, where most "
    "side effects have zero mentions. Gaps persist under normalisation."
)
add_two_figs("tfidf_comparison.png", "side_effects_normalized_rate.png",
             "TF-IDF comparison", "Side-effect rates per 1k tokens", width=2.4)

doc.add_heading("3.3 Classification", level=2)
doc.add_paragraph(
    f"Na\u00efve Bayes (\u03b1=0.1): 97.6% CV accuracy. KNN (k=7): 96.7%. "
    f"Normalised-track accuracy nearly identical (97.6%/96.4%), confirming separability "
    f"is vocabulary-driven. Top public features: anyone, day, lb, feeling; "
    f"top media: cnn, healthline, medscape, nbc."
)

doc.add_heading("3.4 Topic Modelling", level=2)
doc.add_paragraph(
    "LDA (5 topics/corpus): public topics centre on dosing, insurance, weight progress, "
    "compounding; media on clinical trials, FDA regulation, Medicare access. "
    "K-means (k=2) purity: 84.1%\u2014strong unsupervised separation."
)

doc.add_heading("3.5 Aspect-Based Sentiment", level=2)
add_table(
    ["Aspect", "Public", "Media", "Gap", "Sig."],
    [
        ["Mental health", "\u22120.10", "\u22120.74", "0.63", "p=.040*"],
        ["Access", "+0.20", "\u22120.35", "0.55", "p<.001*"],
        ["Efficacy", "+0.14", "\u22120.33", "0.46", "p<.001*"],
        ["Cost", "+0.16", "\u22120.02", "0.18", "p=.017*"],
        ["Side effects", "\u22120.03", "\u22120.20", "0.16", "p=.666"],
    ]
)
doc.add_paragraph(
    "\nMedia is dramatically more negative on mental health (\u22120.74 vs. \u22120.10), "
    "likely driven by alarming headlines. Discrepancy classifier accuracy: 83.7%."
)

doc.add_heading("3.6 Temporal Lead-Lag", level=2)
doc.add_paragraph(
    "Granger causality (max lag 4 weeks, 45 obs.): no significant relationship "
    "(Media\u2192Public p=0.37; Public\u2192Media p=0.30). Cross-correlation peaks "
    "at lag \u22127 weeks (r=0.31). The streams respond to shared events rather "
    "than driving each other."
)
add_two_figs("aspect_sentiment_comparison.png", "crosscorrelation.png",
             "Aspect sentiment", "Cross-correlation", width=2.4)

# ══════════════════════════════════════════════════════════
# 4. DISCUSSION & CONCLUSION
# ══════════════════════════════════════════════════════════
doc.add_heading("4. Discussion & Conclusion", level=1)
doc.add_paragraph(
    "The sentiment gap aligns with prior research: media defaults to cautionary framing "
    "while patients share mixed lived experiences. The 97.6% classification accuracy\u2014"
    "stable under normalisation\u2014confirms fundamentally different registers. "
    "The side-effect coverage gap has real-world implications: patients relying on media "
    "may underestimate everyday symptom burden, while providers could monitor patient "
    "communities for emerging signals. The mental-health discrepancy is striking: media "
    "amplifies alarm while patients report only mild negativity."
)
doc.add_paragraph(
    "Limitations: hybrid mode used 100% snippet fallback; VADER may miss medical-domain "
    "nuance; 45 weekly observations limit Granger power; Reddit/WebMD are not "
    "representative of all patients."
)
doc.add_paragraph(
    "Public and media discourse on GLP-1 drugs diverge significantly in sentiment, "
    "vocabulary, and thematic emphasis. These differences are robust across methods and "
    "length normalisation. Patients relying solely on media may form an incomplete picture; "
    "healthcare providers should monitor both discourse streams."
)

rep = doc.add_paragraph()
run = rep.add_run(
    "Reproducibility: python project_cli.py run-analysis --media-text-mode hybrid"
)
run.italic = True
run.font.size = Pt(8)

# ── Save ──────────────────────────────────────────────────
out = ROOT / "report" / "report.docx"
doc.save(str(out))
print(f"DOCX saved to {out}")
